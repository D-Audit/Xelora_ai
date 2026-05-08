"""
ingestion/document_ingest.py
Unstructured Data Intelligence: turns PDFs, images (via OCR), and Word
documents into plain text and, where possible, tabular rows - so the
Excel agent can pull data in from something that isn't already a
spreadsheet.

Honesty note on scope: this covers text extraction and simple table
detection (PDF tables via pdfplumber, DOCX tables via python-docx).
Handwritten notes depend entirely on OCR quality (pytesseract/Tesseract) -
messy handwriting will extract poorly or not at all; that's a real
limitation of OCR generally, not something this module can work around.
Emails and "mixed data sources" aren't handled as a distinct format here -
route an exported .txt/.eml body through extract_text_file, or a PDF/
screenshot of one through the PDF/image paths below.
"""

import os


def extract_pdf(path: str) -> dict:
    """Returns {'text': str, 'tables': [[[cell,...],...], ...]} - one
    tables entry per detected table, each a list of rows."""
    import pdfplumber

    full_text, tables = [], []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                full_text.append(page_text)
            for table in page.extract_tables():
                tables.append(table)

    return {"text": "\n".join(full_text), "tables": tables, "source_type": "pdf"}


def extract_image(path: str) -> dict:
    """OCR text extraction from an image. Quality depends entirely on
    image clarity and Tesseract being installed on this machine."""
    import pytesseract
    from PIL import Image

    image = Image.open(path)
    text = pytesseract.image_to_string(image)
    return {"text": text, "tables": [], "source_type": "image"}


def extract_docx(path: str) -> dict:
    """Extracts paragraph text plus any tables (as row-lists) from a
    Word document."""
    import docx

    document = docx.Document(path)
    full_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

    tables = []
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(rows)

    return {"text": full_text, "tables": tables, "source_type": "docx"}


def extract_text_file(path: str) -> dict:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return {"text": text, "tables": [], "source_type": "text"}


_EXTRACTORS = {
    ".pdf": extract_pdf,
    ".png": extract_image, ".jpg": extract_image, ".jpeg": extract_image,
    ".bmp": extract_image, ".tiff": extract_image,
    ".docx": extract_docx,
    ".txt": extract_text_file, ".csv": extract_text_file, ".eml": extract_text_file,
}


class UnsupportedFileType(Exception):
    pass


def extract_document(path: str) -> dict:
    """Single entry point: picks the right extractor by file extension
    and returns a unified {'text', 'tables', 'source_type'} result."""
    ext = os.path.splitext(path)[1].lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileType(
            f"No extractor for '{ext}' files. Supported: {sorted(_EXTRACTORS.keys())}"
        )
    return extractor(path)
